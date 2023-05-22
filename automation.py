import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from natsort import natsorted
import time

doc_name = input("Enter the name of the word file: ")
# Start the timer
start_time = time.time()

# Print the task
print("Creating the word file...")

# Get the current directory
directory = os.getcwd()

# Get the list of image filenames in the directory
image_filenames = [filename for filename in os.listdir(directory) if filename.endswith(('.jpg', '.jpeg', '.png')) and filename.lower().endswith('.jpg')]

# Sort the filenames using natural sorting
image_filenames = natsorted(image_filenames)

# Create a new Word document
doc = Document()

# Adjust the document margin
sections = doc.sections
for section in sections:
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

# Create a table with 2 columns and specify the column widths
doc.add_table(rows=1, cols=2)
table = doc.tables[0]
table.columns[0].width = Inches(4)
table.columns[1].width = Inches(4)

# Add table borders and adjust cell alignment
table.style = 'Table Grid'
for row in table.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# Add the text "IMAGE TO BE ANNOTATED" in bold to the first column's first row
cell = table.cell(0, 0)
cell.text = "IMAGE TO BE ANNOTATED"
paragraph = cell.paragraphs[0]
run = paragraph.runs[0]
run.bold = True
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add the text "PROPERTIES" in bold to the second column's first row
cell = table.cell(0, 1)
cell.text = "PROPERTIES"
paragraph = cell.paragraphs[0]
run = paragraph.runs[0]
run.bold = True
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Iterate over the image filenames and add them to the table
for filename in image_filenames:
    # Add a new row to the table
    row = table.add_row().cells

    # Add the image to the first column
    filename_cell = row[0]
    paragraph = filename_cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(os.path.join(directory, filename), width=Inches(4.7))

    # Add the properties section to the second column
    properties_cell = row[1]
    properties_paragraph = properties_cell.paragraphs[0]
    properties_paragraph.add_run(f"Filename: ")
    properties_paragraph.add_run(os.path.splitext(filename)[0]).bold = True
    properties_paragraph.add_run("\n\nPosition of Mandibular Canal\n").bold = True
    properties_paragraph.add_run("(Lingual, apical, buccal): ")
    properties_paragraph.add_run("__________")
    properties_paragraph.add_run("\n\nInterruption of corticalization\n").bold = True
    properties_paragraph.add_run("(Please write")
    properties_paragraph.add_run("'Y'").bold = True  
    properties_paragraph.add_run("if there is an interruption "
                                 "of corticalization.\nOtherwise write ")
    properties_paragraph.add_run("'N'): ").bold = True
    properties_paragraph.add_run("__________")
    properties_paragraph.add_run("\n\nNote: please put “")
    properties_paragraph.add_run("NA").bold = True
    properties_paragraph.add_run("” if the MC is not present or cannot be seen.")


# Save the Word document
doc.save(f'{doc_name}.docx')

def clear_terminal():
    """Clear the terminal screen."""
    os.system('cls' if os.name == 'nt' else 'clear')

# End the timer
end_time = time.time()

# Calculate the runtime
runtime = end_time - start_time
runtime = round(runtime, 2)

clear_terminal()
print(f"\nImagine doing it in {runtime} seconds using this vs doing it manually ><")
print("Word file created successfully!")
